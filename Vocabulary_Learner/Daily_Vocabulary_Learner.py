"""
Daily Vocabulary Learner
========================

USAGE (examples)
----------------
# Run Spanish vocabulary test (day of date = 14)
python word_automation.py
    -> Input: Spanish
    -> Input: 14
    -> Input: 2   # test only level 1 and above

# Run French vocabulary test (day of date = 28), include both tests
python word_automation.py
    -> Input: French
    -> Input: 28
    -> Input: both

# Run low-level test only (French, day of date = 7)
python word_automation.py
    -> Input: French
    -> Input: 7
    -> Input: 1   # test only level 0 and below


Notes
-----
- Prompts user for **language**: 'French' or 'Spanish'.
- Loads corresponding vocabulary DOCX:
    * "Spanish Minimum Core Vocabulary - Copy.docx"
    * "French Minimum Core Vocabulary - Copy.docx"
- Prompts for **date** (int). Used to compute review schedule:
    * 28-day (monthly), 14-day, 7-day, 2-day intervals.
- Levels represent word familiarity (0–6) with specific colors:
    * L0: Bright Green
    * L1: Yellow
    * L2: Gray
    * L3: Pink
    * L4: Turquoise
    * L5: Red
    * L6: Blue
- Tests:
    * `low_test()` → runs through level 0/unknown words.
    * `test()` → runs through scheduled review words (levels 1–6).
- Words can be promoted (advance familiarity) or demoted (manual input).
- Updates word cell colors in the DOCX file and saves changes.

Outputs
-------
* Updates in-place:
    - "Spanish Minimum Core Vocabulary - Copy.docx"
    - "French Minimum Core Vocabulary - Copy.docx"
* Printed test sessions in console:
    - Word prompts, user answers, correctness checks.
    - List of demoted words displayed at end.
"""

import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

is_spanish = input("French or Spanish? ") == "Spanish"
if is_spanish:
    doc = docx.Document("Spanish Minimum Core Vocabulary - Copy.docx")
else:
    doc = docx.Document("French Minimum Core Vocabulary - Copy.docx")
level_0_cells = []
level_1_cells = []
level_2_cells = []
level_3_cells = []
level_4_cells = []
level_5_cells = []
level_6_cells = []

english_columns = [2, 4, 6, 8]


limit = len(doc.tables[0].rows) - 1

one_month_level = [0]
two_week_level = [0]
one_week_level_2 = [0]
one_week_level_1 = [0]
every_other_day_2 = [0]
every_other_day_1 = [0]

init_levels = [
    one_month_level,
    two_week_level,
    one_week_level_2,
    every_other_day_2,
]
numbers_for_levels = [28, 14, 7, 2]

date = int(input("Please enter the date: "))

for z in range(4):
    if date % numbers_for_levels[z] == 0:
        init_levels[z][0] = numbers_for_levels[z]
    else:
        init_levels[z][0] = date % numbers_for_levels[z]
    current = 0
    while init_levels[z][current] + numbers_for_levels[z] <= limit:
        init_levels[z].append(init_levels[z][current] + numbers_for_levels[z])
        current += 1
one_week_level_1 = one_week_level_2
every_other_day_1 = every_other_day_2
levels = [
    one_month_level,
    two_week_level,
    one_week_level_2,
    one_week_level_1,
    every_other_day_2,
    every_other_day_1,
]

demoted_words = []


def length_of_test():
    j = 0
    level_differentiator = 7
    for level in levels:
        level_differentiator -= 1
        for day in level:
            for column in english_columns:

                cell_to_be_checked = doc.tables[0].cell(day, column - 1)
                if cell_to_be_checked.text == "":
                    pass
                elif level_differentiator == 1 and cell_to_be_checked.paragraphs[
                    0
                ].runs[0].font.color.rgb != (0xFF, 0xFF, 0x00):
                    pass
                elif level_differentiator == 2 and cell_to_be_checked.paragraphs[
                    0
                ].runs[0].font.color.rgb != (0xA6, 0xA6, 0xA6):
                    pass
                elif level_differentiator == 3 and cell_to_be_checked.paragraphs[
                    0
                ].runs[0].font.color.rgb != (0xFF, 0x00, 0xFF):
                    pass
                elif level_differentiator == 4 and cell_to_be_checked.paragraphs[
                    0
                ].runs[0].font.color.rgb != (0x00, 0xFF, 0xFF):
                    pass
                elif level_differentiator == 5 and cell_to_be_checked.paragraphs[
                    0
                ].runs[0].font.color.rgb != (0xFF, 0x00, 0x00):
                    pass
                elif level_differentiator == 6 and cell_to_be_checked.paragraphs[
                    0
                ].runs[0].font.color.rgb != (0x00, 0x00, 0xFF):
                    pass
                else:
                    j += 1
    return j


def test():
    j = 0
    print("")
    print("Start of test:")
    level_differentiator = 7
    for level in levels:
        level_differentiator -= 1
        for day in level:
            for column in english_columns:
                cell_to_be_checked = doc.tables[0].cell(day, column - 1)
                cell_to_be_printed = doc.tables[0].cell(day, column)
                if cell_to_be_checked.text == "":
                    pass
                elif level_differentiator == 1 and cell_to_be_checked.paragraphs[
                    0
                ].runs[0].font.color.rgb != (0xFF, 0xFF, 0x00):
                    pass
                elif level_differentiator == 2 and cell_to_be_checked.paragraphs[
                    0
                ].runs[0].font.color.rgb != (0xA6, 0xA6, 0xA6):
                    pass
                elif level_differentiator == 3 and cell_to_be_checked.paragraphs[
                    0
                ].runs[0].font.color.rgb != (0xFF, 0x00, 0xFF):
                    pass
                elif level_differentiator == 4 and cell_to_be_checked.paragraphs[
                    0
                ].runs[0].font.color.rgb != (0x00, 0xFF, 0xFF):
                    pass
                elif level_differentiator == 5 and cell_to_be_checked.paragraphs[
                    0
                ].runs[0].font.color.rgb != (0xFF, 0x00, 0x00):
                    pass
                elif level_differentiator == 6 and cell_to_be_checked.paragraphs[
                    0
                ].runs[0].font.color.rgb != (0x00, 0x00, 0xFF):
                    pass
                else:
                    j += 1
                    if column - 1 == 1:
                        word_type = "(noun)"
                    elif column - 1 == 3:
                        word_type = "(adjective)"
                    elif column - 1 == 5:
                        word_type = "(verb)"
                    elif column - 1 == 7:
                        word_type = "(others)"
                    if j < 10:
                        print(
                            "0"
                            + str(j)
                            + "."
                            + " "
                            + cell_to_be_printed.text
                            + " "
                            + word_type
                            + " "
                            + str(level_differentiator)
                        )
                    else:
                        print(
                            str(j)
                            + "."
                            + " "
                            + cell_to_be_printed.text
                            + " "
                            + word_type
                            + " "
                            + str(level_differentiator)
                        )
                    input("    Answer: ")
                    print("    " + cell_to_be_checked.text)
                    right_or_wrong = input("    Right or wrong? ") == ""
                    if right_or_wrong:
                        if level_differentiator == 6:
                            print("    At maximum familiarity")
                        elif level_differentiator == 5:
                            level_6_cells.append(cell_to_be_checked)
                        elif level_differentiator == 4:
                            level_5_cells.append(cell_to_be_checked)
                        elif level_differentiator == 3:
                            level_4_cells.append(cell_to_be_checked)
                        elif level_differentiator == 2:
                            level_3_cells.append(cell_to_be_checked)
                        elif level_differentiator == 1:
                            level_2_cells.append(cell_to_be_checked)
                        elif level_differentiator == 0:
                            level_1_cells.append(cell_to_be_checked)
                        print("")
                    else:
                        demoted_level = int(
                            input("    Which level shall this word be demoted to? ")
                        )
                        if demoted_level == 6:
                            level_6_cells.append(cell_to_be_checked)
                        if demoted_level == 5:
                            level_5_cells.append(cell_to_be_checked)
                        if demoted_level == 4:
                            level_4_cells.append(cell_to_be_checked)
                        if demoted_level == 3:
                            level_3_cells.append(cell_to_be_checked)
                        if demoted_level == 2:
                            level_2_cells.append(cell_to_be_checked)
                        if demoted_level == 1:
                            level_1_cells.append(cell_to_be_checked)
                        if demoted_level == 0:
                            level_0_cells.append(cell_to_be_checked)
                        demoted_words.append(cell_to_be_checked.text)
                        print("")
    print("End of test")


def low_test():
    print("Start of low level test: ")
    for row in range(limit):
        for column in [1, 3, 5, 7]:
            if doc.tables[0].cell(row, column).text == "":
                pass
            elif doc.tables[0].cell(row, column).paragraphs[0].runs[
                0
            ].font.color.rgb == (0x00, 0xFF, 0x00):
                if column == 1:
                    word_type = "(noun)"
                elif column == 3:
                    word_type = "(adjective)"
                elif column == 5:
                    word_type = "(verb)"
                elif column == 7:
                    word_type = "(other)"
                print(doc.tables[0].cell(row, column + 1).text + " " + word_type)
                input("Answer: ")
                print(doc.tables[0].cell(row, column).text)
                if input("Where shall this word be placed? ") == "":
                    level_1_cells.append(doc.tables[0].cell(row, column))
                    print("")
            """elif doc.tables[0].cell(row, column).paragraphs[0].runs[0].font.highlight_color == WD_COLOR_INDEX.GRAY_25:
                if column == 1:
                    word_type = '(noun)'
                elif column == 3:
                    word_type = '(adjective)'
                elif column == 5:
                    word_type = '(verb)'
                elif column == 7:
                    word_type = '(other)'
                print(doc.tables[0].cell(row, column + 1).text + ' ' + word_type)
                input('Answer: ')
                print(doc.tables[0].cell(row, column).text)
                if input('Where shall this word be placed? ') == '':
                    level_0_cells.append(doc.tables[0].cell(row, column))"""
    print("End of low level test. ")


def change_level():
    for i in range(len(level_0_cells)):
        for run in level_0_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
    for i in range(len(level_1_cells)):
        for run in level_1_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0x00)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    for i in range(len(level_2_cells)):
        for run in level_2_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xA6, 0xA6, 0xA6)
            run.font.highlight_color = WD_COLOR_INDEX.GRAY_50
    for i in range(len(level_3_cells)):
        for run in level_3_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0xFF)
            run.font.highlight_color = WD_COLOR_INDEX.PINK
    for i in range(len(level_4_cells)):
        for run in level_4_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0x00, 0xFF, 0xFF)
            run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
    for i in range(len(level_5_cells)):
        for run in level_5_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.font.highlight_color = WD_COLOR_INDEX.RED
    for i in range(len(level_6_cells)):
        for run in level_6_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
            run.font.highlight_color = WD_COLOR_INDEX.BLUE


wanted_test = input("Do you want to do level 0 and below, level 1 and above, or both? ")
if wanted_test == "1":
    low_test()
elif wanted_test == "2":
    print(f"{length_of_test()} words to be tested.")
    test()
else:
    print(f"{length_of_test()} words to be tested.")
    test()
    low_test()
print("")
print(demoted_words)
print("")
if input("Please confirm colour change (do not type yes): ") == "":
    change_level()
if is_spanish:
    doc.save("Spanish Minimum Core Vocabulary - Copy.docx")
else:
    doc.save("French Minimum Core Vocabulary - Copy.docx")
